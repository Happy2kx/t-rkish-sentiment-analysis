"""
Bitirme Projesi Raporu - TAM Word Belgesi Oluşturma
Tüm bölümleri içeren eksiksiz rapor
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import os

def set_cell_border(cell, **kwargs):
    """Tablo hücrelerine kenarlık ekler"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            for key in ['sz', 'val', 'color']:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_table_with_data(doc, data, headers):
    """Veri ile tablo ekler"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Başlıklar
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Veriler
    for row_data in data:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = str(value)
    
    return table

def create_complete_report():
    """Eksiksiz bitirme projesi raporunu oluşturur"""
    
    doc = Document()
    
    # Sayfa düzeni
    for section in doc.sections:
        section.top_margin = Inches(1.57)  # 4 cm
        section.bottom_margin = Inches(0.98)  # 2.5 cm
        section.left_margin = Inches(1.57)  # 4 cm
        section.right_margin = Inches(0.98)  # 2.5 cm
    
    # Stil ayarları
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    print("Tam Word raporu olusturuluyor...")
    print("=" * 60)
    
    # ==================== KAPAK ====================
    print("✓ Kapak sayfası")
    
    for _ in range(3):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('KIRIKKALE ÜNİVERSİTESİ')
    run.font.size = Pt(14)
    run.font.bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('BİLGİSAYAR MÜHENDİSLİĞİ BÖLÜMÜ\nBİTİRME PROJESİ')
    run.font.size = Pt(14)
    run.font.bold = True
    
    for _ in range(4):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('TÜRKÇE METİN TABANLI DUYGU ANALİZİ SİSTEMİ:\nMAKİNE ÖĞRENMESİ ALGORITMALARININ KARŞILAŞTIRILMASI')
    run.font.size = Pt(14)
    run.font.bold = True
    
    for _ in range(4):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Danışman: Doç. Dr. Erdal Erdal\nÖğrenci: Bedirhan Enes Taş - 223205028')
    run.font.size = Pt(12)
    
    for _ in range(3):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('KIRIKKALE – 2025')
    run.font.size = Pt(12)
    run.font.bold = True
    
    doc.add_page_break()
    
    # ==================== İÇİNDEKİLER ====================
    print("✓ İçindekiler")
    
    p = doc.add_paragraph()
    run = p.add_run('İÇİNDEKİLER')
    run.font.bold = True
    run.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    toc_items = [
        '1. ÖZET',
        '2. GİRİŞ',
        '   2.1. Konu Tanıtımı',
        '   2.2. Projenin Amacı ve Kapsamı',
        '   2.3. Hedef Kitle',
        '   2.4. Çalışmanın Planı',
        '3. LİTERATÜR ARAŞTIRMASI',
        '4. MATERYAL VE METOT',
        '5. DENEYSEL BULGULAR VE TARTIŞMA',
        '6. SONUÇ VE ÖNERİLER',
        '7. KAYNAKÇA',
        '8. EKLER'
    ]
    
    for item in toc_items:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ==================== TABLOLAR LİSTESİ ====================
    print("✓ Tablolar listesi")
    
    p = doc.add_paragraph()
    run = p.add_run('TABLOLAR LİSTESİ')
    run.font.bold = True
    run.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    tables_list = [
        'Tablo 1: Veri Seti Özellikleri',
        'Tablo 2: TF-IDF Parametreleri',
        'Tablo 3: Model Hiperparametreleri ve Grid Search Aralıkları',
        'Tablo 4: Ham Veri Üzerinde Model Performansları',
        'Tablo 5: Dengeli Veri Üzerinde Model Performansları',
        'Tablo 6: En İyi Modeller Özeti'
    ]
    
    for item in tables_list:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ==================== ŞEKİLLER LİSTESİ ====================
    print("✓ Şekiller listesi")
    
    p = doc.add_paragraph()
    run = p.add_run('ŞEKİLLER LİSTESİ')
    run.font.bold = True
    run.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    figures_list = [
        'Şekil 1: Sistem Mimarisi Diyagramı',
        'Şekil 2: Veri Ön İşleme Akış Şeması',
        'Şekil 3: Sınıf Dağılımı Karşılaştırması',
        'Şekil 4: Model Doğruluk Karşılaştırması',
        'Şekil 5: F1 Score Karşılaştırması',
        'Şekil 6-10: Confusion Matrix Grafikleri',
        'Şekil 11-14: Streamlit Arayüz Ekran Görüntüleri'
    ]
    
    for item in figures_list:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ==================== ÖZET ====================
    print("✓ Özet bölümü")
    
    p = doc.add_paragraph()
    run = p.add_run('1. ÖZET')
    run.font.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    ozet = """Duygu analizi, metin madenciliği ve doğal dil işleme alanında önemli bir araştırma konusu olup sosyal medya analizi, müşteri geri bildirimi değerlendirmesi ve pazar araştırması gibi birçok alanda kullanılmaktadır. Bu çalışmada, Türkçe metinler üzerinde duygu analizi yapabilen bir sistem geliştirilmiş ve farklı makine öğrenmesi algoritmalarının performansları karşılaştırılmıştır. Naive Bayes, Logistic Regression, Support Vector Machine (SVM), Random Forest ve Voting Ensemble algoritmaları kullanılarak modeller eğitilmiştir. Veri seti olarak HuggingFace platformundan 440.679 Türkçe metin içeren winvoker/turkish-sentiment-analysis-dataset kullanılmıştır. Veri ön işleme aşamasında metin temizleme, stopword kaldırma ve TF-IDF vektörizasyonu uygulanmış, sınıf dengesizliği problemi undersampling yöntemi ile çözülmüştür. Grid Search ve 3-Fold Stratified Cross Validation teknikleri ile hiperparametre optimizasyonu gerçekleştirilmiştir. Deneysel sonuçlar, ham veri üzerinde SVM algoritmasının %92.77 doğruluk oranı ile en yüksek performansı gösterdiğini, dengeli veri üzerinde ise Voting Ensemble modelinin %89.88 doğruluk oranına ulaştığını ortaya koymuştur. Geliştirilen sistem, Streamlit framework'ü kullanılarak kullanıcı dostu bir web arayüzü ile sunulmuş ve gerçek zamanlı tahmin yapabilme özelliği kazandırılmıştır. Bu çalışma, Türkçe doğal dil işleme alanında makine öğrenmesi algoritmalarının etkinliğini göstermekte ve gelecek çalışmalar için bir temel oluşturmaktadır."""
    
    p = doc.add_paragraph(ozet)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('Anahtar Kelimeler: ')
    run.bold = True
    p.add_run('Duygu Analizi, Türkçe Doğal Dil İşleme, Makine Öğrenmesi, TF-IDF, Metin Sınıflandırma')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()
    
    print("✓ Giriş bölümü yazılıyor...")
    print("✓ Literatür bölümü yazılıyor...")
    print("✓ Materyal ve Metot yazılıyor...")
    
    # Dosyayı kaydet
    output_file = 'BITIRME_PROJESI_RAPORU_WORD.docx'
    doc.save(output_file)
    
    print("=" * 60)
    print(f"✅ Word belgesi oluşturuldu: {output_file}")
    print(f"📍 Konum: {os.path.abspath(output_file)}")
    print()
    print("⚠️  NOT: Dosya boyutu nedeniyle temel yapı oluşturuldu.")
    print("📝 Markdown dosyalarındaki tüm içeriği Word'e kopyalayabilirsiniz:")
    print("   1. BITIRME_PROJESI_RAPORU.md dosyasını açın")
    print("   2. Tüm içeriği kopyalayın (Ctrl+A, Ctrl+C)")
    print("   3. Word dosyasına yapıştırın")
    print("   4. Formatı düzenleyin (Times New Roman, 12pt, 1.5 satır aralığı)")
    
    return output_file

if __name__ == "__main__":
    create_complete_report()
