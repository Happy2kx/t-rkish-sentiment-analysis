# -*- coding: utf-8 -*-
"""
Bitirme Projesi Raporu - Word Formatinda
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_word_report():
    doc = Document()
    
    # Sayfa duzeni
    for section in doc.sections:
        section.top_margin = Inches(1.57)
        section.bottom_margin = Inches(0.98)
        section.left_margin = Inches(1.57)
        section.right_margin = Inches(0.98)
    
    # Normal stil
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    print("Word raporu olusturuluyor...")
    
    # KAPAK
    for _ in range(3):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('KIRIKKALE UNIVERSITESI')
    run.font.size = Pt(14)
    run.font.bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('BILGISAYAR MUHENDISLIGI BOLUMU')
    run.font.size = Pt(14)
    run.font.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('BITIRME PROJESI')
    run.font.size = Pt(14)
    run.font.bold = True
    
    for _ in range(4):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('TURKCE METIN TABANLI DUYGU ANALIZI SISTEMI:')
    run.font.size = Pt(14)
    run.font.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('MAKINE OGRENMESI ALGORITMALARININ KARSILASTIRILMASI')
    run.font.size = Pt(14)
    run.font.bold = True
    
    for _ in range(4):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Danisman: Doc. Dr. Erdal Erdal')
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Ogrenci: Bedirhan Enes Tas - 223205028')
    run.font.size = Pt(12)
    
    for _ in range(3):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('KIRIKKALE - 2025')
    run.font.size = Pt(12)
    run.font.bold = True
    
    doc.add_page_break()
    
    # Dosyayi kaydet
    output_file = 'BITIRME_PROJESI_RAPORU.docx'
    doc.save(output_file)
    
    print("Basarili! Word dosyasi olusturuldu: " + output_file)
    print("Konum: " + os.path.abspath(output_file))
    print()
    print("ONEMLI NOT:")
    print("Word dosyasinin temel yapisi olusturuldu.")
    print("Simdi Markdown dosyalarindaki icerigi Word'e kopyalayin:")
    print("1. BITIRME_PROJESI_RAPORU.md dosyasini acin")
    print("2. Tum icerigi kopyalayin")
    print("3. Word dosyasina yapiştirin")
    print("4. Formati duzenleyin")
    
    return output_file

if __name__ == "__main__":
    create_word_report()
