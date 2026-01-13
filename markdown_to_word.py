# -*- coding: utf-8 -*-
"""
Markdown'dan Word'e Donusturucu
RAPOR_TAM.md dosyasini duzgun formatlanmis Word'e cevirir
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def create_formatted_word():
    print("Word belgesi olusturuluyor...")
    
    # Yeni belge
    doc = Document()
    
    # Sayfa duzeni
    section = doc.sections[0]
    section.top_margin = Inches(1.57)  # 4 cm
    section.bottom_margin = Inches(0.98)  # 2.5 cm
    section.left_margin = Inches(1.57)  # 4 cm
    section.right_margin = Inches(0.98)  # 2.5 cm
    
    # Normal stil
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    print("Markdown dosyasi okunuyor...")
    
    # Markdown oku
    with open('RAPOR_TAM.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Satirlara ayir
    lines = content.split('\n')
    
    print(f"Toplam {len(lines)} satir okundu")
    print("Word'e donusturuluyor...")
    
    current_paragraph = ''
    in_table = False
    table_lines = []
    
    for i, line in enumerate(lines):
        line = line.rstrip()
        
        # Bos satirlari atla
        if not line or line == '---':
            if current_paragraph:
                p = doc.add_paragraph(current_paragraph)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing = 1.5
                current_paragraph = ''
            continue
        
        # Baslik 1 (# )
        if line.startswith('# '):
            if current_paragraph:
                p = doc.add_paragraph(current_paragraph)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                current_paragraph = ''
            
            text = line[2:].strip()
            p = doc.add_paragraph()
            run = p.add_run(text.upper())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            
        # Baslik 2 (## )
        elif line.startswith('## '):
            if current_paragraph:
                p = doc.add_paragraph(current_paragraph)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                current_paragraph = ''
            
            text = line[3:].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(5)
            
        # Baslik 3 (### )
        elif line.startswith('### '):
            if current_paragraph:
                p = doc.add_paragraph(current_paragraph)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                current_paragraph = ''
            
            text = line[4:].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
        
        # Kalin metin (**text**)
        elif line.startswith('**') and line.endswith('**') and len(line) > 4:
            if current_paragraph:
                p = doc.add_paragraph(current_paragraph)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                current_paragraph = ''
            
            text = line[2:-2]
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Tablo
        elif '|' in line and not line.startswith('```'):
            # Tablo satirlarini topla
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
        
        # Kod blogu
        elif line.startswith('```'):
            if current_paragraph:
                p = doc.add_paragraph(current_paragraph)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                current_paragraph = ''
            # Kod bloklarini atla (simdilik)
            continue
        
        # Normal metin
        else:
            if in_table:
                # Tablo bitti, olustur
                in_table = False
                # Simdilik tabloyu atla (karmasik)
                table_lines = []
            
            # Metni birlestir
            current_paragraph += line + ' '
        
        # Ilerleme
        if i % 100 == 0:
            print(f"Islenen satir: {i}/{len(lines)}")
    
    # Son paragraf
    if current_paragraph:
        p = doc.add_paragraph(current_paragraph)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
    
    # Kaydet
    output_file = 'BITIRME_RAPORU_FORMATTED.docx'
    doc.save(output_file)
    
    print(f"\nBasarili! Dosya kaydedildi: {output_file}")
    print(f"Konum: {output_file}")
    
    return output_file

if __name__ == "__main__":
    create_formatted_word()
